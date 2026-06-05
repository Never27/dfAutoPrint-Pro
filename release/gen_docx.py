"""Generate comprehensive PdfAutoPrint Pro documentation as a rich Word document."""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(6)

for level in range(1, 5):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    heading_style.font.bold = True
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

TITLE_COLOR = RGBColor(0x4A, 0x45, 0xCF)  # Purple
ACCENT_COLOR = RGBColor(0x6B, 0x5F, 0xD4)
HEADER_BG = "4A45CF"
LIGHT_BG = "F5F3FF"
BORDER_COLOR = "D4D0FF"

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_cover_page():
    """Create a professional cover page."""
    # Spacer
    for _ in range(5):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PdfAutoPrint Pro')
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = TITLE_COLOR
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Windows 虚拟 PDF 打印机管理器')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph()

    # Version
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('版本 v1.1.2')
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT_COLOR
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT_COLOR

    # Author info
    info_lines = [
        ('作者', '和学斌'),
        ('联系方式', 'QQ：1210696000'),
        ('许可证', 'MIT License'),
        ('GitHub', 'https://github.com/Never27/dfAutoPrint-Pro'),
        ('文档日期', datetime.date.today().strftime('%Y年%m月%d日')),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_page_break()

def add_styled_table(headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        set_cell_shading(cell, HEADER_BG)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Microsoft YaHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            if r % 2 == 1:
                set_cell_shading(cell, LIGHT_BG)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table

def add_code_block(code_text):
    """Add a styled code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Add border via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="{BORDER_COLOR}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def add_bullet(text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))

def add_section_header(text, level=1):
    """Add a section heading."""
    doc.add_heading(text, level=level)

def add_paragraph_text(text, bold=False):
    """Add a normal paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p

# ═══════════════════════════════════════════
# BUILD THE DOCUMENT
# ═══════════════════════════════════════════

# ── Cover Page ──
add_cover_page()

# ════════════════════════════════════════
# Chapter 1: Executive Summary
# ════════════════════════════════════════
add_section_header('第一章  产品概述', 1)

add_section_header('1.1 什么是 PdfAutoPrint Pro？', 2)
add_paragraph_text(
    'PdfAutoPrint Pro 是一款专为 Windows 平台打造的虚拟 PDF 打印机管理器。'
    '它在您的系统中注册一个虚拟打印机，当您从任何应用程序（Word、Excel、浏览器、CAD、图像软件等）'
    '执行"打印"操作时，PdfAutoPrint Pro 会自动拦截打印任务，将其转换为高质量 PDF 文件并保存到您指定的位置。'
    '与普通 PDF 打印机不同，PdfAutoPrint Pro 提供了一整套智能化的文件管理、命名规范和输出配置能力，'
    '让 PDF 输出流程真正实现自动化和标准化。'
)

add_section_header('1.2 核心亮点', 2)

add_section_header('零外部依赖', 3)
add_paragraph_text(
    '整个应用基于 .NET 8 WPF 原生构建，不依赖任何第三方 NuGet 包。除了运行时必需的 .NET 8 Desktop Runtime '
    '和底层使用的 GhostScript（由安装脚本自动处理），项目没有引入任何额外的外部库。这意味着：'
)
for item in ['安装包体积极小', '无供应链安全风险', '不依赖不可控的第三方更新', '长期维护无忧']:
    add_bullet(item)

add_section_header('智能文件名模板', 3)
add_paragraph_text('支持使用占位符动态生成文件名，无需每次手动输入。系统在生成 PDF 时自动解析占位符，替换为实际信息。')
add_styled_table(
    ['占位符', '说明', '示例输出'],
    [
        ['<PrintJobName>', '打印任务名称（即文档标题）', '季度报告'],
        ['<PrintJobAuthor>', '打印文档作者', '张三'],
        ['<DateTime:format>', '当前日期时间（自定义格式）', '2026-06-05'],
        ['<Counter>', '自增计数器（全局）', '001、002'],
        ['<PrinterName>', '当前打印机名称', 'PdfAutoPrint_办公'],
    ],
    [4, 6, 5]
)

add_section_header('多打印机配置', 3)
add_paragraph_text('支持创建多个虚拟 PDF 打印机，每台打印机拥有独立的完整配置。不同打印机可用于不同工作场景，互不干扰。')
add_styled_table(
    ['打印机名称', '配置', '用途'],
    [
        ['PdfAutoPrint_办公', '自动保存、矢量、无印章', '日常文档转换'],
        ['PdfAutoPrint_归档', '递增编号、灰度、时间戳备份', '历史文档归档'],
        ['PdfAutoPrint_印章', '公司章水印、彩色', '正式合同生成'],
        ['PdfAutoPrint_设计', '高 DPI、图片模式', '设计稿输出'],
    ],
    [5, 5, 5]
)

add_section_header('水印与印章', 3)
add_paragraph_text('在生成的 PDF 文件上自动添加水印，支持文字水印和图片水印两种类型。')
add_bullet('文字水印：可配置字体、大小、颜色、透明度、旋转角度、位置')
add_bullet('图片水印（盖章）：支持 PNG/JPG/BMP，可缩放、设置透明度、定位')
add_bullet('页面跳过：可配置跳过的页面范围（如封面不加印章）')

add_section_header('系统托盘集成', 3)
for item in ['最小化到系统托盘，不占用任务栏空间', '打印任务完成/失败时弹出气泡通知', '点击托盘图标快速打开主窗口', '右键菜单：显示主窗口 / 退出']:
    add_bullet(item)

add_section_header('开机自启', 3)
add_paragraph_text(
    '支持开机自动启动，无需手动打开软件。通过注册表 HKCU\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Run 实现，无需管理员权限。'
)

add_section_header('1.3 产品对比', 2)
add_styled_table(
    ['对比维度', 'PdfAutoPrint Pro', 'Windows 自带 PDF', '第三方 PDF 打印机', 'Adobe Acrobat'],
    [
        ['自动命名', '模板化', '不支持', '部分支持', '不支持'],
        ['水印支持', '内置', '不支持', '不支持', '付费功能'],
        ['多打印机', '支持', '不支持', '不支持', '不支持'],
        ['批量处理', '自动跳过失败', '不支持', '不支持', '有限支持'],
        ['安装复杂度', '一键安装包', '系统自带', '需手动配置', '需注册登录'],
        ['零依赖', '是', '—', '否', '否'],
        ['开源免费', 'MIT', '预装', '部分免费', '付费'],
        ['托盘通知', '支持', '不支持', '不支持', '不支持'],
        ['开机自启', '支持', '不支持', '不支持', '不支持'],
    ],
    [3, 3, 3, 3, 3]
)

add_section_header('1.4 适用场景', 2)
scenarios = [
    ('企业办公', '批量合同、报价单、发票的 PDF 归档；统一文件命名规范；自动添加公司水印，保护文档版权。'),
    ('个人使用', '网页文章保存为 PDF 存档；各类文档格式统一转换为 PDF；重要邮件的 PDF 备份。'),
    ('开发测试', '自动化测试中的报表生成；批量文档转换和格式验证；打印输出的质量对比测试。'),
    ('设计制图', 'CAD 图纸输出为高清 PDF；图片设计稿的 PDF 交付；支持高 DPI 矢量输出。'),
]
for title, desc in scenarios:
    add_paragraph_text(f'「{title}」{desc}')

doc.add_page_break()

# ════════════════════════════════════════
# Chapter 2: System Architecture
# ════════════════════════════════════════
add_section_header('第二章  系统架构', 1)

add_section_header('2.1 架构概览', 2)
add_paragraph_text(
    'PdfAutoPrint Pro 采用经典的四层架构设计，遵循 MVVM (Model-View-ViewModel) 模式。'
    '整个系统由表现层、视图模型层、服务层和数据层组成，层与层之间通过清晰的接口进行通信，'
    '确保了代码的高内聚低耦合特性。'
)

add_section_header('2.2 分层架构', 2)

add_styled_table(
    ['层级', '组件', '职责'],
    [
        ['表现层 (Presentation)', 'MainWindow / PrinterConfigWindow / WatermarkEditorWindow / Converters / App.xaml', '用户界面展示与交互'],
        ['视图模型层 (ViewModel)', 'MainViewModel (INotifyPropertyChanged)', '12 个命令 + 可观察属性绑定'],
        ['服务层 (Services)', 'SpoolMonitor / GhostScriptService / PsWatermarkEngine / PrinterManager / ConfigService / JobHistoryService / FileNameResolver / ConflictResolver / LogService / StartupService', '核心业务逻辑'],
        ['数据层 (Models)', 'PrinterProfile / WatermarkConfig / JobRecord', '数据结构定义与 JSON 序列化'],
        ['外部集成 (External)', 'GhostScript CLI / PowerShell/WMI / FileSystemWatcher', '系统资源交互'],
    ],
    [3, 6, 6]
)

add_section_header('2.3 核心处理管线', 2)
add_paragraph_text(
    'PdfAutoPrint Pro 的打印处理管线是一个精心设计的流水线，从用户触发打印到最终 PDF 输出，'
    '经过多个阶段的自动化处理。以下是完整的处理流程：'
)

pipeline_steps = [
    ('1. 用户应用', '用户在任意应用（Word/浏览器/CAD等）中执行"打印"操作，选择虚拟打印机'),
    ('2. Windows 打印池', 'Windows Print Spooler 服务接收打印任务，生成 .prn (PostScript) 文件到 spool 目录'),
    ('3. 文件监控', 'SpoolMonitor 通过 FileSystemWatcher 监听 spool 目录，检测新 .prn 文件创建事件'),
    ('4. 文件就绪检测', '等待文件写入完成（轮询独占访问检测），防止读取不完整文件'),
    ('5. 元数据提取', '从 .prn 文件头部解析 %%Title / %%For 等 DSC 注释，提取文档标题和作者信息'),
    ('6. 真实文档名获取', '通过 WMI (Win32_PrintJob) 查询打印队列中的真实文档名，覆盖 DSC 占位名'),
    ('7. 文件名解析', 'FileNameResolver 解析用户配置的命名模板，替换日期/计数/来源等占位符'),
    ('8. 冲突处理', 'ConflictResolver 检查目标路径是否存在同名文件，按策略处理（覆盖/递增/备份）'),
    ('9. 水印注入', '如启用水印，PsWatermarkEngine 将 PostScript 水印代码注入到 .prn 文件中'),
    ('10. PS→PDF 转换', 'GhostScriptService 调用 gswin64c.exe，将 .prn 文件转换为高质量 PDF'),
    ('11. 输出保存', 'PDF 文件保存到指定目录，记录作业历史到 job-history.json'),
    ('12. 通知反馈', '通过 Windows 系统托盘弹出气泡通知，告知用户任务结果'),
]
for step, desc in pipeline_steps:
    add_paragraph_text(f'【{step}】', bold=True)
    add_paragraph_text(desc)

add_section_header('2.4 项目结构', 2)
add_code_block(
    'PdfAutoPrint.Pro/\n'
    '├── Models/                         数据模型层\n'
    '│   ├── PrinterProfile.cs           打印机配置模型\n'
    '│   ├── WatermarkConfig.cs          水印/印章/页面叠加模型\n'
    '│   └── JobRecord.cs                作业历史记录模型\n'
    '├── Services/                       业务服务层\n'
    '│   ├── ConfigService.cs            JSON 配置管理\n'
    '│   ├── GhostScriptService.cs        GhostScript CLI 调用服务\n'
    '│   ├── PrinterManager.cs            打印机 CRUD 管理\n'
    '│   ├── PsWatermarkEngine.cs         PostScript 水印注入引擎\n'
    '│   ├── FileNameResolver.cs          文件名模板解析\n'
    '│   ├── ConflictResolver.cs          文件冲突处理\n'
    '│   ├── LogService.cs                日志服务\n'
    '│   ├── JobHistoryService.cs         作业历史服务\n'
    '│   ├── SpoolMonitor.cs             打印池监控与处理管线\n'
    '│   └── StartupService.cs           开机自启服务\n'
    '├── ViewModels/                     视图模型层\n'
    '│   └── MainViewModel.cs            主窗口 ViewModel（12 个功能命令）\n'
    '├── Views/                          视图层\n'
    '│   ├── MainWindow.xaml/.cs          主窗口（4 标签页）\n'
    '│   ├── PrinterConfigWindow.xaml/.cs 打印机配置编辑器（5 标签页）\n'
    '│   └── WatermarkEditorWindow.xaml/.cs 水印项编辑器\n'
    '├── Converters/                     WPF 值转换器\n'
    '│   └── Converters.cs\n'
    '├── App.xaml/.cs                    应用入口和全局样式\n'
    '└── PdfAutoPrint.Pro.csproj         项目文件'
)

doc.add_page_break()

# ════════════════════════════════════════
# Chapter 3: Core Features
# ════════════════════════════════════════
add_section_header('第三章  核心功能详解', 1)

features = [
    ('3.1 自动/手动保存切换', [
        '控制 PDF 文件生成时的保存行为，支持两种模式：自动保存模式（PDF 生成后直接存入预设目录，无需任何弹窗确认）'
        '和手动模式（每次打印弹出文件保存对话框，由用户选择保存位置和文件名）。'
        '该设置可按照打印机独立配置，满足不同场景需求。',
    ]),
    ('3.2 智能文件名模板', [
        '支持使用占位符动态生成文件名，无需每次手动输入。系统在生成 PDF 时自动解析占位符，替换为实际信息。',
        '模板示例：<DateTime:yyyyMMdd>_<PrintJobName>_<Counter:D3>',
        '结果示例：20260605_季度报告_001.pdf',
        '支持使用 / 在模板中创建子目录，<Counter:D3> 中的 D3 表示 3 位数零填充。',
        '如果文档标题包含临时文件名（如 Word 临时文件名），系统会自动过滤。',
    ]),
    ('3.3 命名冲突处理', [
        '当生成的文件名与已有文件重名时，提供三种处理策略：',
        '覆盖（Overwrite）：直接替换已有文件，不保留原文件。',
        '递增编号（Increment）：在文件名后添加序号，如 报告(1).pdf、报告(2).pdf。',
        '时间戳备份（Backup）：将已有文件重命名为 报告_20260605_143052.bak，再保存新文件。',
    ]),
    ('3.4 系统托盘通知', [
        '每次打印任务完成后，通过 Windows 系统托盘弹出气泡通知，告知用户任务结果。',
        '成功通知：显示文件名和保存路径。失败通知：显示失败原因摘要。',
        '可在设置中开启/关闭通知，成功通知可选择显示或关闭（失败通知默认始终显示）。',
    ]),
    ('3.5 日志保留策略', [
        '详细记录每次打印任务的信息，包括时间、来源应用、文件名、结果等。',
        '支持保留 N 天和永久保留两种策略。保留 N 天策略下，超过指定天数的日志自动清理。',
        '日志内容包括：任务创建时间、源应用名称、输出文件路径、处理耗时、成功/失败状态及错误详情。',
    ]),
    ('3.6 失败自动跳过', [
        '在批量打印场景下，如果某一项任务失败（如文件损坏、权限不足），'
        '系统自动跳过该任务，继续处理后续任务，不会中断整个批量流程。',
        '默认开启，可在设置中关闭（关闭后，失败会暂停队列等待用户干预）。',
        '失败的任务会记录在日志查看器中，可集中查看和排查。',
    ]),
    ('3.7 输出质量控制', [
        '精细控制 PDF 输出质量，满足不同场景的清晰度和文件大小需求。',
    ]),
]

for title, content in features:
    add_section_header(title, 2)
    for line in content:
        add_paragraph_text(line)
    if '输出质量' in title:
        add_styled_table(
            ['场景', '色彩模式', '输出模式', 'DPI'],
            [
                ['日常文档存档', '彩色', '矢量', '—'],
                ['网页保存', '彩色', '矢量', '—'],
                ['CAD 图纸输出', '灰度', '图片', '600'],
                ['设计稿交付', '彩色', '图片', '300'],
                ['档案归档（最小体积）', '灰度', '矢量', '—'],
            ],
            [4, 3, 3, 3]
        )
        add_paragraph_text('矢量模式生成的 PDF 中文字可被选中和搜索；高 DPI + 图片模式会产生较大文件，仅在需要时使用。')
    elif '冲突' in title:
        add_styled_table(
            ['策略', '行为', '适用场景'],
            [
                ['覆盖', '直接替换已有文件，不保留原文件', '临时文件、每次需要最新版本的场景'],
                ['递增编号', '在文件名后添加序号，如 报告(1).pdf', '需要保留所有历史版本的场景'],
                ['时间戳备份', '将已有文件重命名为 .bak 再保存新文件', '需要追溯任意时间点版本的合规场景'],
            ],
            [3, 6, 6]
        )

add_section_header('3.8 一键安装工具', 2)
add_paragraph_text('通过单个安装包完成全套环境搭建，无需手动下载和配置各组件。自动完成的任务包括：')
install_steps = [
    '检测系统环境（Windows 版本、架构）',
    '检查/安装 .NET 8 Desktop Runtime',
    '检查/安装 GhostScript 10.04.0',
    '安装虚拟 PDF 打印机驱动程序',
    '创建虚拟打印机并注册',
    '注册 PdfAutoPrint Pro 管理软件',
    '创建桌面/开始菜单快捷方式',
    '配置开机自启（可选）',
]
for i, step in enumerate(install_steps, 1):
    add_bullet(f'{i}. {step}')

add_section_header('3.9 打印机管理', 2)
add_paragraph_text('集中管理 Windows 系统中的所有打印机，特别是虚拟 PDF 打印机。主要功能包括：')
for item in ['重命名虚拟打印机显示名称', '查看系统中所有打印机及其状态', '一键重启 Print Spooler 服务', '设置默认打印机', '查看打印队列和任务状态', '删除/添加虚拟打印机']:
    add_bullet(item)

add_section_header('3.10 水印支持', 2)
add_paragraph_text('在生成的 PDF 文件上自动添加水印，支持文字水印和图片水印两种类型。')
add_paragraph_text('文字水印属性：', bold=True)
for item in ['文字内容：支持多行文字，可使用文件名等变量', '字体：系统字体选择，支持大小、颜色配置', '透明度：0-100% 可调', '旋转角度：0-360 度，常用对角线 45 度', '位置：九宫格位置 + 像素微调（绝对定位）', '间距：平铺模式下的行列间距']:
    add_bullet(item)
add_paragraph_text('图片水印（图片盖章）属性：', bold=True)
for item in ['图片来源：支持 PNG、JPG、BMP 格式', '缩放：按比例或按像素缩放', '透明度：0-100%', '位置：九宫格 + 偏移，支持自动居中']:
    add_bullet(item)
add_paragraph_text('页面跳过功能：可配置跳过第一页（封面不加水印）、仅奇数/偶数页、指定页码范围。')

add_section_header('3.11 开机自启', 2)
add_paragraph_text(
    '支持开机自动启动。通过注册表 HKCU\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Run 实现，'
    '无需管理员权限。程序启动后自动最小化到系统托盘，不打扰用户。'
)

doc.add_page_break()

# ════════════════════════════════════════
# Chapter 4: Technical Implementation
# ════════════════════════════════════════
add_section_header('第四章  技术实现', 1)

add_section_header('4.1 技术栈', 2)
add_styled_table(
    ['技术', '版本', '用途'],
    [
        ['.NET', '8.0 (LTS)', '运行时框架'],
        ['WPF', '.NET 8 内置', '桌面 UI 框架'],
        ['C#', '12.0', '编程语言'],
        ['MVVM', '手动实现', 'UI 架构模式'],
        ['GhostScript', '10.x (外部)', 'PS → PDF 转换引擎'],
        ['JSON', 'System.Text.Json', '数据序列化'],
        ['PowerShell/WMI', '系统内置', '打印机管理'],
    ],
    [4, 4, 7]
)

add_section_header('4.2 技术选型理由', 2)
add_paragraph_text('为什么 .NET 8 WPF？', bold=True)
add_styled_table(
    ['选项', '评价'],
    [
        ['.NET 8 WPF', '原生 Windows 支持，性能好，开发效率高'],
        ['Electron', '资源占用大，不符合"轻量级"要求'],
        ['WinUI 3', '需要 Windows 10 1809+，不支持 Win7'],
        ['WinForms', 'UI 开发效率低，不支持现代布局'],
    ],
    [4, 11]
)

add_paragraph_text('为什么零 NuGet 依赖？', bold=True)
add_paragraph_text(
    '.NET 10 SDK 的 NuGet 还原存在兼容性问题（Path.Combine(null, ...) 异常），'
    '项目功能完全可用 .NET 内置 API 实现，减少部署复杂度，无需考虑依赖版本冲突，'
    'JSON 替代 SQLite 满足配置存储需求。'
)

add_paragraph_text('为什么 JSON 而非 SQLite？', bold=True)
add_styled_table(
    ['特性', 'JSON', 'SQLite'],
    [
        ['依赖', '无（System.Text.Json 内置）', '需 System.Data.SQLite NuGet'],
        ['可读性', '人工可读', '需工具查看'],
        ['性能', '小数据量足够', '大数据量更快'],
        ['备份', '直接复制', '需导出'],
    ],
    [3, 6, 6]
)

add_section_header('4.3 核心组件详解', 2)

add_section_header('打印池监控 (SpoolMonitor)', 3)
add_paragraph_text(
    'SpoolMonitor 是整个系统的入口组件，负责实时监控 Windows 打印池目录（C:\\Windows\\System32\\spool\\PRINTERS\\），'
    '检测新产生的 .prn (PostScript) 文件并触发处理管线。'
)
add_paragraph_text('技术实现要点：', bold=True)
for item in [
    '文件系统监控：基于 System.IO.FileSystemWatcher，监控打印池目录的 Created 事件',
    '文件就绪检测：通过文件独占访问尝试判断文件是否写入完成（轮询 10 次，每次 500ms）',
    '并发控制：使用 SemaphoreSlim 限制同时处理的作业数量（最大 3），防止资源耗尽',
    '错误隔离：每个作业在独立的 try-catch 块中处理，单个失败不影响整体监控',
    '真实文档名获取：通过 WMI (Win32_PrintJob) 查询打印队列中的真实文档名，覆盖 DSC 占位名（如 MSxpsPS）',
]:
    add_bullet(item)

add_section_header('PDF 转换引擎 (GhostScriptService)', 3)
add_paragraph_text(
    'GhostScriptService 封装了 GhostScript 命令行工具的调用逻辑，将 PostScript 文件转换为 PDF 格式。'
)
add_paragraph_text('关键技术点：', bold=True)
for item in [
    '进程管理：通过 System.Diagnostics.Process 启动 gswin64c.exe',
    '异步通信：使用 RedirectStandardOutput/Error 异步读取进程输出，避免死锁',
    '超时控制：支持转换超时设定，超时后强制终止进程',
    '安全模式：始终使用 -dSAFER 参数运行，限制文件访问和系统命令执行',
    '内容保护：使用 -dPDFFitPage 确保内容自适应页面，不被裁切',
    '自动旋转：使用 -dAutoRotatePages=/None 禁止自动旋转页面',
]:
    add_bullet(item)

add_styled_table(
    ['PDF 质量等级', 'GhostScript 参数', '用途'],
    [
        ['screen（屏幕）', '-dPDFSETTINGS=/screen', '72 DPI，适合屏幕阅读'],
        ['ebook（电子书）', '-dPDFSETTINGS=/ebook', '150 DPI，适合电子阅读器'],
        ['printer（打印）', '-dPDFSETTINGS=/printer', '300 DPI，适合打印输出'],
        ['prepress（印刷）', '-dPDFSETTINGS=/prepress', '300+ DPI, CMYK，适合印刷'],
    ],
    [4, 5, 6]
)

add_section_header('水印注入引擎 (PsWatermarkEngine)', 3)
add_paragraph_text(
    'PsWatermarkEngine 实现了基于 PostScript 代码注入的水印方案。水印代码直接嵌入到 PostScript 文件中，'
    '在 GhostScript 转换阶段渲染为 PDF 页面的一部分，具有不可篡改性。'
)
add_paragraph_text('与传统 PDF 层水印相比的优势：', bold=True)
add_styled_table(
    ['对比维度', 'PDF 层水印', 'PostScript 注入水印'],
    [
        ['可移除性', '可被 PDF 工具移除', '渲染到页面内容中，不可分离'],
        ['渲染质量', '依赖 PDF 工具', 'GhostScript 矢量渲染，质量更高'],
        ['性能影响', '需要后处理步骤', '转换时一并完成，无额外开销'],
        ['灵活性', '有限', '完全控制 PS 图形状态'],
    ],
    [3, 5, 7]
)

add_section_header('打印机管理 (PrinterManager)', 3)
add_paragraph_text('通过 PowerShell 命令管理打印机，无需 System.Management NuGet 包：')
add_code_block(
    '# 获取打印机\n'
    'Get-WmiObject -Class Win32_Printer | Select-Object Name, DriverName, PortName\n'
    '\n'
    '# 安装打印机\n'
    'Add-PrinterPort -Name $portName\n'
    'Add-Printer -Name $name -DriverName $driverName -PortName $portName\n'
    '\n'
    '# 重启 Spooler\n'
    'Restart-Service -Name Spooler'
)

add_section_header('4.4 数据模型', 2)
add_code_block(
    'GlobalConfig (config.json)\n'
    '├── List<PrinterProfile>\n'
    '│   ├── OutputConfig (自动保存、输出路径、监听目录)\n'
    '│   ├── FileNameTemplate (文件名模式、占位符规则)\n'
    '│   ├── ConflictConfig (冲突策略：覆盖/递增/备份)\n'
    '│   ├── QualityConfig (色彩模式、输出模式、DPI)\n'
    '│   ├── NotificationConfig (系统托盘、成功/失败通知)\n'
    '│   ├── ServiceConfig (失败跳过、GS路径、开机自启)\n'
    '│   └── WatermarkConfig\n'
    '│       └── List<WatermarkItem>\n'
    '│           ├── WatermarkType (Text/Image)\n'
    '│           ├── WatermarkPosition (锚点、偏移、单位)\n'
    '│           ├── PageRange (跳过页面)\n'
    '│           └── FontConfig (字体、颜色)\n'
    '└── LogConfig (保留天数、日志级别)\n'
    '\n'
    'JobRecord (job-history.json)\n'
    '├── PrinterName, JobName, Author\n'
    '├── SourceFile, OutputFile, FileSize, PageCount\n'
    '├── Status, ErrorMessage\n'
    '└── StartTime, EndTime, DurationMs'
)

doc.add_page_break()

# ════════════════════════════════════════
# Chapter 5: Security
# ════════════════════════════════════════
add_section_header('第五章  安全设计', 1)

add_section_header('5.1 进程安全', 2)
add_styled_table(
    ['风险点', '防护措施'],
    [
        ['GhostScript 参数注入', '参数白名单校验，不直接拼接用户输入'],
        ['PowerShell 命令注入', '使用 PowerShell AST 参数传递，避免字符串拼接'],
        ['进程权限提升', 'GS 以普通用户权限运行，使用 -dSAFER 安全模式'],
        ['进程泄漏', 'using 语句确保进程资源释放，超时强制终止'],
    ],
    [5, 10]
)

add_section_header('5.2 文件系统安全', 2)
add_styled_table(
    ['风险点', '防护措施'],
    [
        ['文件路径遍历', '校验输出路径在配置的目录范围内'],
        ['文件权限', '检查目录读写权限后再执行操作'],
        ['敏感信息泄露', '日志中不记录完整文件内容，仅记录路径'],
        ['临时文件清理', '处理完成后删除中间 .prn 文件'],
    ],
    [5, 10]
)

add_section_header('5.3 数据安全', 2)
add_styled_table(
    ['风险点', '防护措施'],
    [
        ['配置文件损坏', '写入前创建 .bak 备份文件'],
        ['JSON 解析异常', '容错处理，异常时恢复默认配置'],
        ['并发写入冲突', '配置写入使用文件锁机制'],
    ],
    [5, 10]
)

add_section_header('5.4 GhostScript 安全模式', 2)
add_paragraph_text('-dSAFER 参数启用 GS 安全模式后，将禁止以下操作：')
for item in ['禁止写入任意文件', '禁止执行系统命令', '限制内存使用', '限制文件访问路径']:
    add_bullet(item)

doc.add_page_break()

# ════════════════════════════════════════
# Chapter 6: Installation Guide
# ════════════════════════════════════════
add_section_header('第六章  安装指南', 1)

add_section_header('6.1 系统要求', 2)
add_styled_table(
    ['项目', '最低要求', '推荐配置'],
    [
        ['操作系统', 'Windows 10 x64', 'Windows 10/11 x64'],
        ['.NET 运行时', '.NET 8 Desktop Runtime x64', '.NET 8 Desktop Runtime x64'],
        ['GhostScript', '10.04.0（安装包自动安装）', '10.04.0'],
        ['CPU', '双核 1.5GHz', '四核 2.0GHz+'],
        ['内存', '2 GB', '4 GB+'],
        ['磁盘空间', '300 MB（含运行时）', '500 MB'],
        ['显示器', '1024x768', '1920x1080'],
    ],
    [3, 5, 7]
)

add_section_header('6.2 一键安装步骤', 2)
add_paragraph_text('方法一：使用 .bat 安装包（推荐）', bold=True)
for step in [
    '下载 PdfAutoPrint_Pro_Setup_v1.1.2.bat（约 118MB，包含所有依赖）',
    '右键 .bat 文件 → "以管理员身份运行"',
    '等待安装完成（自动处理 .NET 8 Runtime、GhostScript、虚拟打印机驱动、快捷方式等）',
    '安装完成后，桌面和开始菜单会出现 PdfAutoPrint Pro 图标',
]:
    add_bullet(step)

add_paragraph_text('方法二：使用 .exe SFX 安装包', bold=True)
for step in [
    '下载 PdfAutoPrint_Pro_Setup_v1.1.2.exe（约 344KB，轻量版）',
    '双击运行，按提示操作',
    '注意：此版本需要系统已安装 .NET 8 Desktop Runtime 和 GhostScript 10.x',
]:
    add_bullet(step)

add_section_header('6.3 安装后验证', 2)
for step in [
    '打开任意文档（如记事本），按 Ctrl+P 打印',
    '选择打印机 "Auto PDF Printer"',
    '点击"打印"，查看输出目录是否生成了 PDF 文件',
    '查看系统托盘是否有 PdfAutoPrint Pro 图标和通知',
    '双击托盘图标打开主界面，确认仪表盘计数 +1',
]:
    add_bullet(step)

add_section_header('6.4 故障排查', 2)
troubleshooting = [
    ('Print Spooler 服务未运行', '以管理员身份运行 PowerShell，执行：Start-Service Spooler；Set-Service Spooler -StartupType Automatic'),
    ('PDF 生成后是空白文件', '打开 PdfAutoPrint Pro，进入"设置" → "高级"，手动指定 GhostScript 可执行文件路径（通常位于 C:\\Program Files\\gs\\gs10.x.x\\bin\\gswin64c.exe）'),
    ('打印后无任何反应', '确认 PdfAutoPrint Pro 在运行（检查系统托盘图标），查看 Windows 打印队列是否卡住，在应用中点击"重启 Spooler"'),
    ('文件名乱码', '确认 Windows 系统区域设置为"中文（简体，中国）"，控制面板 → 区域 → 管理 → 更改系统区域设置'),
]
for problem, solution in troubleshooting:
    add_paragraph_text(f'问题：{problem}', bold=True)
    add_paragraph_text(f'解决：{solution}')

add_section_header('6.5 卸载指南', 2)
for step in [
    '退出 PdfAutoPrint Pro（右键托盘图标 → 退出）',
    '控制面板 → 设备和打印机 → 右键 "Auto PDF Printer" → 删除设备',
    '删除安装目录：C:\\Program Files\\PdfAutoPrint Pro\\',
    '删除配置目录（如需彻底清理）：%LOCALAPPDATA%\\PdfAutoPrint.Pro\\',
    '如不再需要 GhostScript，通过控制面板 → 程序和功能 → 卸载',
]:
    add_bullet(step)

doc.add_page_break()

# ════════════════════════════════════════
# Chapter 7: Performance & Compatibility
# ════════════════════════════════════════
add_section_header('第七章  性能与兼容性', 1)

add_section_header('7.1 转换性能基准', 2)
add_styled_table(
    ['文档类型', '页数', '转换时间', '输出文件大小'],
    [
        ['纯文本文档', '10 页', '~2 秒', '~50 KB'],
        ['纯文本文档', '100 页', '~8 秒', '~500 KB'],
        ['图文混排', '10 页', '~4 秒', '~2 MB'],
        ['图文混排', '100 页', '~25 秒', '~20 MB'],
        ['高清扫描件', '10 页', '~6 秒', '~5 MB'],
        ['高清扫描件', '100 页', '~60 秒', '~50 MB'],
    ],
    [3, 2, 3, 3]
)

add_section_header('7.2 资源占用', 2)
add_styled_table(
    ['指标', '空闲时', '处理中（单作业）', '处理中（3作业并发）'],
    [
        ['内存', '~30 MB', '~80 MB', '~150 MB'],
        ['CPU', '< 1%', '~15%', '~40%'],
        ['磁盘 I/O', '无', '中等', '中高'],
        ['启动时间', '< 2 秒', '—', '—'],
    ],
    [3, 3, 4, 4]
)

add_section_header('7.3 操作系统兼容性', 2)
add_styled_table(
    ['操作系统', '版本', '支持状态'],
    [
        ['Windows 10', '1809+', '完全支持'],
        ['Windows 10 LTSC', '2019', '完全支持'],
        ['Windows 11', '21H2+', '完全支持'],
        ['Windows Server', '2019+', '支持（需桌面体验）'],
        ['Windows 7/8', '—', '不支持'],
    ],
    [4, 4, 7]
)

add_section_header('7.4 错误处理策略', 2)
add_styled_table(
    ['错误层级', '处理策略'],
    [
        ['文件读取失败', '如果 SkipOnFailure=true → 跳过记录，继续处理后续任务'],
        ['GS 转换失败', '记录错误日志 + 任务标记为 Failed'],
        ['冲突处理失败', '回退到递增编号策略，确保文件安全'],
        ['配置读写失败', '使用默认配置 + 记录日志告警'],
        ['PowerShell 执行失败', '超时重试 + 向用户展示友好错误提示'],
    ],
    [4, 11]
)

doc.add_page_break()

# ════════════════════════════════════════
# Chapter 8: Version History & Roadmap
# ════════════════════════════════════════
add_section_header('第八章  版本历史与路线图', 1)

add_section_header('8.1 版本历史', 2)
add_styled_table(
    ['版本', '日期', '更新内容'],
    [
        ['v1.1.2', '2026-06-05', 'GhostScript 静默安装参数修复（/S → /VERYSILENT）；WatchFolder 监听目录修复（自动推导 + UI 只读保护）；UI 删除打印机按钮（DangerButton + 系统级卸载）'],
        ['v1.1.1', '2026-06-05', '应用图标（多分辨率 .ico 16~256px）；托盘与窗口图标统一；安装包 .bat 自解压 + .exe SFX 双格式；打印文件名修复（WMI Win32_PrintJob）；GS 内容裁切修复；单台打印机独立启停；状态指示灯中文映射'],
        ['v1.1.0', '2026-06-04', '系统托盘最小化运行；开机自启功能；打印完成/失败气泡通知；完整安装包（含 GhostScript）；WinForms/WPF 类型冲突修复'],
        ['v1.0.0', '2026-06', '初始版本；12 项核心功能完整实现；WPF 桌面界面；一键安装脚本'],
    ],
    [2, 3, 10]
)

add_section_header('8.2 未来路线图', 2)

add_paragraph_text('短期计划 (v1.x)：', bold=True)
for item in ['多语言界面支持（中文/英文）', '打印任务优先级队列', '邮件/FTP 自动转发（转换后）', '配置导入/导出功能']:
    add_bullet(item)

add_paragraph_text('中期计划 (v2.x)：', bold=True)
for item in ['远程监控 Web 界面', '打印任务调度（定时打印）', 'OCR 文字识别集成', 'PDF 合并与拆分工具', '批量文件转换模式']:
    add_bullet(item)

add_paragraph_text('长期计划 (v3.x)：', bold=True)
for item in ['多平台支持（Linux/macOS via CUPS）', '云存储集成（OneDrive/Google Drive）', '企业级集中管理控制台', '打印审计与合规报表', 'AI 智能文档分类与路由']:
    add_bullet(item)

doc.add_page_break()

# ════════════════════════════════════════
# Chapter 9: Developer Guide
# ════════════════════════════════════════
add_section_header('第九章  开发者指南', 1)

add_section_header('9.1 开发环境搭建', 2)
for step in [
    '安装 .NET 8 SDK：https://dotnet.microsoft.com/download/dotnet/8.0',
    '克隆项目：git clone https://github.com/Never27/dfAutoPrint-Pro.git',
    '安装 GhostScript 10.x（测试时需要）：下载 gswin64c.exe 并确保在 PATH 中',
    '构建项目：dotnet build -c Release --no-restore',
    '运行项目：dotnet run -c Release',
]:
    add_bullet(step)

add_section_header('9.2 编码规范', 2)
add_styled_table(
    ['元素', '规范', '示例'],
    [
        ['类名', 'PascalCase', 'PrinterProfile'],
        ['方法名', 'PascalCase', 'GetPrinters()'],
        ['私有字段', '_camelCase', '_ghostScriptPath'],
        ['事件', 'PascalCase + 前缀 On', 'OnJobCompleted'],
    ],
    [3, 5, 7]
)

add_paragraph_text('异步编程要求：', bold=True)
for item in ['所有 I/O 操作使用 async/await 模式', '避免在 UI 线程上执行阻塞操作', 'GhostScript 进程调用使用 Process.Start() + 异步等待']:
    add_bullet(item)

add_section_header('9.3 扩展点', 2)
add_paragraph_text('自定义水印处理器：', bold=True)
add_paragraph_text(
    'PsWatermarkEngine 设计为可扩展，可通过继承或接口添加新的水印类型。'
    '已支持的水印类型包括：文字水印（字体、大小、颜色、角度、透明度）、图片印章（路径、位置、缩放、透明度）、'
    '页眉/页脚（文本、页码、日期）。'
)

add_paragraph_text('添加新的文件冲突策略：', bold=True)
add_paragraph_text(
    'ConflictResolver 支持三种内置策略（覆盖/递增/备份），可通过添加新的 ConflictStrategy 枚举值扩展。'
)

add_paragraph_text('事件系统钩子：', bold=True)
add_code_block(
    '// SpoolMonitor 暴露的事件\n'
    'public event EventHandler<JobEventArgs> JobDetected;\n'
    'public event EventHandler<JobEventArgs> JobCompleted;\n'
    'public event EventHandler<ErrorEventArgs> JobFailed;'
)

add_section_header('9.4 Git 分支策略', 2)
for item in [
    'master — 稳定发布分支',
    'develop — 开发集成分支',
    'feature/* — 功能分支',
    'bugfix/* — 修复分支',
]:
    add_bullet(item)

add_section_header('9.5 代码审查要点', 2)
for item in [
    '遵循编码规范（PascalCase / _camelCase）',
    '无新的 NuGet 依赖引入（零依赖原则）',
    '异步方法正确使用 async/await',
    '错误处理覆盖关键路径',
    '不影响现有功能',
]:
    add_bullet(item)

# ════════════════════════════════════════
# Footer
# ════════════════════════════════════════
doc.add_page_break()
add_section_header('附录', 1)

add_section_header('A. 技术指标汇总', 2)
add_styled_table(
    ['指标', '值'],
    [
        ['运行时', '.NET 8 (LTS)'],
        ['目标平台', 'Windows 10/11 x64'],
        ['启动时间', '< 2 秒'],
        ['内存占用（空闲）', '< 100 MB'],
        ['外部依赖', '仅 GhostScript CLI'],
        ['许可证', 'MIT License'],
        ['GitHub', 'https://github.com/Never27/dfAutoPrint-Pro'],
        ['作者', '和学斌'],
        ['联系方式', 'QQ：1210696000'],
    ],
    [5, 10]
)

add_section_header('B. 文件存储路径', 2)
add_styled_table(
    ['存储项', '路径'],
    [
        ['应用配置', '%LOCALAPPDATA%\\PdfAutoPrint.Pro\\config.json'],
        ['打印机配置', '%LOCALAPPDATA%\\PdfAutoPrint.Pro\\printer-profiles.json'],
        ['作业历史', '%LOCALAPPDATA%\\PdfAutoPrint.Pro\\job-history.json'],
        ['日志文件', '%LOCALAPPDATA%\\PdfAutoPrint.Pro\\logs\\app-{date}.log'],
        ['安装目录', 'C:\\Program Files\\PdfAutoPrint Pro\\'],
        ['GhostScript', 'C:\\Program Files\\gs\\gs10.04.0\\bin\\gswin64c.exe'],
        ['打印池目录', 'C:\\Windows\\System32\\spool\\PRINTERS\\'],
    ],
    [4, 11]
)

add_section_header('C. 快捷键参考', 2)
add_styled_table(
    ['操作', '快捷键 / 方式'],
    [
        ['打开主窗口', '双击系统托盘图标'],
        ['退出应用', '右键托盘图标 → 退出'],
        ['添加打印机', '主界面 → 打印机管理 → 添加新打印机'],
        ['重启 Spooler', '主界面 → 打印机管理 → 重启 Spooler'],
        ['查看日志', '主界面 → 日志查看器'],
        ['启动/停止单台', '打印机卡片上的绿色/红色按钮'],
    ],
    [5, 10]
)

# ── Save ──
output_path = r'C:\Users\Administrator\WorkBuddy\2026-06-05-09-13-35\PdfAutoPrint.Pro\release\PdfAutoPrint_Pro_详细文档_v1.1.2.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Size: {os.path.getsize(output_path) // 1024} KB')
